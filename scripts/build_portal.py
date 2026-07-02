from __future__ import annotations

import html
import json
import re
import shutil
import unicodedata
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = Path(r"C:\Users\User\Downloads\UnespDataLens-RM_Projeto_Completo_Final_Ajustado.docx")
SOURCE_PDF = Path(r"C:\Users\User\Downloads\UnespDataLens-RM_Projeto_Completo_Final_Ajustado.pdf")
LEXICON_ROOT = Path(r"C:\Users\User\Documents\unespdatalens-projeto-completo\docs")

SITE_TITLE = "UnespDataLens-RM"
SITE_SUBTITLE = "Modelo de Referência para Engenharia Analítica de Dados"
DOWNLOAD_DOCX_NAME = SOURCE_DOCX.name
DOWNLOAD_PDF_NAME = SOURCE_PDF.name

CATEGORY_LABELS = {
    "conceitos": "Conceitos",
    "tecnicas": "Técnicas e Tecnologias",
    "metodos": "Métodos e Algoritmos",
    "artefatos": "Artefatos",
    "metricas": "Métricas",
    "problemas": "Problemas Tratados",
    "aplicacoes": "Aplicações",
    "trilhas": "Trilhas de Estudo",
}

CATEGORY_DESCRIPTIONS = {
    "conceitos": "Fundamentos conceituais usados em todo o modelo de referência.",
    "tecnicas": "Técnicas, bibliotecas, plataformas e tecnologias aplicáveis aos pipelines.",
    "metodos": "Métodos de processamento, validação, integração, explicabilidade e análise.",
    "artefatos": "Entregáveis formais, registros, manifestos, mapas e produtos de dados.",
    "metricas": "Indicadores para medir qualidade, cobertura, rastreabilidade e desempenho.",
    "problemas": f"Riscos, falhas e limitações que o {SITE_TITLE} ajuda a enfrentar.",
    "aplicacoes": "Contextos em que o modelo pode ser instanciado e avaliado.",
    "trilhas": "Percursos orientados para estudar o modelo por objetivo ou perfil.",
}

METHOD_TITLES = {
    "Batch Processing",
    "Change Data Capture",
    "Data Profiling",
    "Data Quality Rules",
    "Deduplicação",
    "Entity Resolution",
    "GAT",
    "GCN",
    "GraphSAGE",
    "Imputação",
    "Ingestão Incremental",
    "LIME",
    "Node2Vec",
    "Normalização",
    "Record Linkage",
    "SHAP",
    "Schema Validation",
    "Streaming",
    "Tratamento de Outliers",
    "Validação de Contratos",
}

EDITORIAL_PREFIXES = (
    "segue ",
    "a seguir",
    "se quiser",
    "posso ",
    "agora vou",
    "agora, vou",
    "perfeito.",
    "perfeito,",
    "ótimo.",
    "ótimo,",
    "excelente.",
    "excelente,",
    "como você",
)


def slugify(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    ascii_value = ascii_value.casefold().replace("×", "x")
    ascii_value = re.sub(r"[^a-z0-9]+", "-", ascii_value)
    return ascii_value.strip("-") or "pagina"


def clean_text(value: str) -> str:
    normalized = " ".join(value.replace("\u00a0", " ").split())
    normalized = re.sub(r"\bUnespDataLens(?!-RM)\b", SITE_TITLE, normalized)
    normalized = normalized.replace(
        "Transformação, Limpeza e Enriquecimento",
        "Transformação, Limpeza, Enriquecimento e Preparação dos Dados",
    )
    normalized = normalized.replace(
        "Tabular, Temporal e Grafos",
        "Tabular, Temporal, Espacial, Textual, Semântica e Grafos",
    )
    return normalized


def strip_numbering(value: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+)*[.)]?\s*", "", value).strip()


def is_editorial(value: str) -> bool:
    folded = clean_text(value).casefold()
    return bool(folded) and folded.startswith(EDITORIAL_PREFIXES)


def iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def paragraph_style(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def is_top_level_title(paragraph: Paragraph) -> bool:
    text = clean_text(paragraph.text)
    style = paragraph_style(paragraph)
    if style in {"Heading 1", "Title"}:
        return True
    return text == "Dimensão Complementar — Data Valuation"


def split_document(doc: Document) -> list[dict]:
    sections: list[dict] = []
    current = None
    for block in iter_blocks(doc):
        if isinstance(block, Paragraph) and is_top_level_title(block):
            title = clean_text(block.text)
            if title == "Dimensão Complementar — Data Valuation":
                current = {"title": title, "blocks": [block]}
                sections.append(current)
                continue
            current = {"title": title, "blocks": [block]}
            sections.append(current)
        elif current is not None:
            current["blocks"].append(block)

    deduplicated = []
    seen_titles = set()
    for section in sections:
        key = section["title"].casefold()
        if key in seen_titles:
            continue
        seen_titles.add(key)
        deduplicated.append(section)
    return deduplicated


def read_frontmatter(markdown_text: str) -> tuple[dict, str]:
    metadata: dict[str, str] = {}
    body = markdown_text
    if markdown_text.startswith("---"):
        parts = markdown_text.split("---", 2)
        if len(parts) == 3:
            for line in parts[1].splitlines():
                if ":" in line and not line.lstrip().startswith("-"):
                    key, value = line.split(":", 1)
                    metadata[key.strip()] = value.strip()
            body = parts[2].lstrip()
    return metadata, body


def load_entries() -> dict[str, list[dict]]:
    categories = ("conceitos", "tecnicas", "artefatos", "metricas", "problemas", "aplicacoes", "trilhas")
    result: dict[str, list[dict]] = {}
    for category in categories:
        entries = []
        folder = LEXICON_ROOT / category
        for source in sorted(folder.glob("*.md")):
            if source.stem == "index":
                continue
            raw = source.read_text(encoding="utf-8")
            metadata, body = read_frontmatter(raw)
            heading = re.search(r"^#\s+(.+)$", body, flags=re.MULTILINE)
            title = clean_text(metadata.get("title") or (heading.group(1) if heading else source.stem))
            entries.append(
                {
                    "title": title,
                    "slug": source.stem,
                    "category": category,
                    "body": body,
                    "url": f"{category}/{source.stem}.html",
                }
            )
        result[category] = entries

    technique_by_title = {entry["title"]: entry for entry in result["tecnicas"]}
    result["metodos"] = []
    for title in sorted(METHOD_TITLES):
        source = technique_by_title.get(title)
        if not source:
            continue
        result["metodos"].append(
            {
                **source,
                "category": "metodos",
                "url": f"metodos/{source['slug']}.html",
            }
        )
    return result


def markdown_excerpt(body: str) -> str:
    text = re.sub(r"!\[[^\]]*]\([^)]+\)", "", body)
    text = re.sub(r"\[([^\]]+)]\([^)]+\)", r"\1", text)
    text = re.sub(r"[#*_`>|-]", " ", text)
    lines = [clean_text(line) for line in text.splitlines() if clean_text(line)]
    for line in lines:
        if len(line) >= 55 and not line.casefold().startswith(("links relacionados", "módulos")):
            return line[:260]
    return (lines[0] if lines else f"Verbete de referência do {SITE_TITLE}.")[:260]


def inline_markdown(value: str) -> str:
    placeholders: list[str] = []

    def save_link(match):
        label = html.escape(clean_text(match.group(1)))
        url = match.group(2).replace(".md", ".html")
        url = re.sub(r"(?:\.\./)?modulos/modulo-(\d{2})-[^#]+\.html", r"../modulos/m\1.html", url)
        url = url.replace(
            "../dimensoes/dados-temporais-eventos-e-tempo-real.html",
            "../modulos/m13.html",
        )
        token = f"\x00LINK{len(placeholders)}\x00"
        placeholders.append(f'<a href="{html.escape(url, quote=True)}">{label}</a>')
        return token

    value = re.sub(r"\[([^\]]+)]\(([^)]+)\)", save_link, value)
    escaped = html.escape(value)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    escaped = re.sub(r"\*([^*]+)\*", r"<em>\1</em>", escaped)
    for index, replacement in enumerate(placeholders):
        escaped = escaped.replace(html.escape(f"\x00LINK{index}\x00"), replacement)
        escaped = escaped.replace(f"\x00LINK{index}\x00", replacement)
    return escaped


def markdown_to_html(body: str) -> str:
    lines = body.splitlines()
    output: list[str] = []
    paragraph: list[str] = []
    list_kind = None

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            text = " ".join(clean_text(line) for line in paragraph)
            if text:
                output.append(f"<p>{inline_markdown(text)}</p>")
            paragraph = []

    def close_list():
        nonlocal list_kind
        if list_kind:
            output.append(f"</{list_kind}>")
            list_kind = None

    first_h1_skipped = False
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            flush_paragraph()
            close_list()
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            close_list()
            level = len(heading.group(1))
            title = clean_text(heading.group(2))
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                continue
            anchor = slugify(title)
            output.append(f'<h{level} id="{anchor}">{inline_markdown(title)}</h{level}>')
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        ordered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if bullet or ordered:
            flush_paragraph()
            desired = "ul" if bullet else "ol"
            if list_kind != desired:
                close_list()
                css_class = "list" if desired == "ul" else "steps"
                output.append(f'<{desired} class="{css_class}">')
                list_kind = desired
            item = bullet.group(1) if bullet else ordered.group(1)
            output.append(f"<li>{inline_markdown(clean_text(item))}</li>")
            continue
        if line.startswith("!!! "):
            flush_paragraph()
            close_list()
            label = line[4:].strip().strip('"')
            output.append(f'<div class="box tip"><strong>{inline_markdown(label)}</strong></div>')
            continue
        paragraph.append(line)
    flush_paragraph()
    close_list()
    return "\n".join(output)


def numbering_formats(doc: Document) -> dict[tuple[int, int], str]:
    formats: dict[tuple[int, int], str] = {}
    try:
        root = doc.part.numbering_part.element
    except Exception:
        return formats
    abstract_ids: dict[int, int] = {}
    for num in root.findall(qn("w:num")):
        num_id = int(num.get(qn("w:numId")))
        abstract = num.find(qn("w:abstractNumId"))
        if abstract is not None:
            abstract_ids[num_id] = int(abstract.get(qn("w:val")))
    abstracts = {}
    for abstract in root.findall(qn("w:abstractNum")):
        abstracts[int(abstract.get(qn("w:abstractNumId")))] = abstract
    for num_id, abstract_id in abstract_ids.items():
        abstract = abstracts.get(abstract_id)
        if abstract is None:
            continue
        for level in abstract.findall(qn("w:lvl")):
            ilvl = int(level.get(qn("w:ilvl")))
            num_fmt = level.find(qn("w:numFmt"))
            if num_fmt is not None:
                formats[(num_id, ilvl)] = num_fmt.get(qn("w:val"))
    return formats


def list_info(paragraph: Paragraph, formats: dict[tuple[int, int], str]) -> tuple[bool, str]:
    ppr = paragraph._p.pPr
    if ppr is None or ppr.numPr is None:
        return False, "ul"
    try:
        num_id = int(ppr.numPr.numId.val)
        ilvl = int(ppr.numPr.ilvl.val) if ppr.numPr.ilvl is not None else 0
        fmt = formats.get((num_id, ilvl), "bullet")
        return True, "ol" if fmt in {"decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman"} else "ul"
    except Exception:
        return True, "ul"


def block_text(block) -> str:
    if isinstance(block, Paragraph):
        return clean_text(block.text)
    if isinstance(block, Table):
        return " ".join(clean_text(cell.text) for row in block.rows for cell in row.cells)
    return ""


def smart_excerpt(value: str, limit: int = 360) -> str:
    value = clean_text(value)
    if len(value) <= limit:
        return value
    candidate = value[:limit]
    sentence_end = max(candidate.rfind("."), candidate.rfind("!"), candidate.rfind("?"))
    if sentence_end >= int(limit * 0.58):
        return candidate[: sentence_end + 1]
    return candidate.rsplit(" ", 1)[0].rstrip(" ,;:") + "…"


def section_summary(section: dict) -> str:
    seek = False
    for block in section["blocks"][1:]:
        if not isinstance(block, Paragraph):
            continue
        text = clean_text(block.text)
        style = paragraph_style(block)
        if not text or is_editorial(text):
            continue
        if style.startswith("Heading"):
            seek = any(word in text.casefold() for word in ("finalidade", "definição", "introdução", "objetivo geral"))
            continue
        if seek and len(text) > 50:
            return smart_excerpt(text)
    for block in section["blocks"][1:]:
        text = block_text(block)
        if len(text) > 60 and not is_editorial(text):
            return smart_excerpt(text)
    return "Conteúdo integral estruturado a partir do documento de referência."


def build_term_targets(entries_by_category: dict[str, list[dict]]) -> list[tuple[str, str, re.Pattern]]:
    priority = ("metodos", "conceitos", "metricas", "artefatos", "problemas", "tecnicas")
    targets: dict[str, tuple[str, str]] = {}
    for category in reversed(priority):
        for entry in entries_by_category[category]:
            title = entry["title"]
            if len(title) < 3:
                continue
            targets[title.casefold()] = (title, entry["url"])
    ordered = sorted(targets.values(), key=lambda item: len(item[0]), reverse=True)
    return [
        (
            title,
            url,
            re.compile(rf"(?<![\w-])({re.escape(title)})(?![\w-])", re.IGNORECASE),
        )
        for title, url in ordered
    ]


def linkify_text(
    value: str,
    page_prefix: str,
    term_targets: list[tuple[str, str, re.Pattern]],
    linked_terms: set[str],
) -> str:
    candidates = []
    for title, url, pattern in term_targets:
        key = title.casefold()
        if key in linked_terms:
            continue
        match = pattern.search(value)
        if not match:
            continue
        candidates.append((match.start(), match.end(), key, url, match.group(1)))

    if not candidates:
        return html.escape(value)

    selected = []
    cursor = 0
    for start, end, key, url, matched_text in sorted(
        candidates, key=lambda item: (item[0], -(item[1] - item[0]))
    ):
        if start < cursor:
            continue
        selected.append((start, end, key, url, matched_text))
        cursor = end

    output = []
    cursor = 0
    for start, end, key, url, matched_text in selected:
        output.append(html.escape(value[cursor:start]))
        output.append(
            f'<a class="term" href="{html.escape(page_prefix + url, quote=True)}">'
            f"{html.escape(matched_text)}</a>"
        )
        linked_terms.add(key)
        cursor = end
    output.append(html.escape(value[cursor:]))
    return "".join(output)


def prepare_heading_anchors(section: dict) -> dict[int, str]:
    result = {}
    used: defaultdict[str, int] = defaultdict(int)
    for block in section["blocks"]:
        if not isinstance(block, Paragraph):
            continue
        style = paragraph_style(block)
        text = clean_text(block.text)
        if not text or not style.startswith("Heading"):
            continue
        base = slugify(strip_numbering(text))
        used[base] += 1
        anchor = base if used[base] == 1 else f"{base}-{used[base]}"
        result[id(block)] = anchor
    return result


def prepare_context_paragraphs(section: dict) -> list[tuple[str, str]]:
    contexts: list[tuple[str, str]] = []
    current_anchor = ""
    for block in section["blocks"]:
        if not isinstance(block, Paragraph):
            continue
        text = clean_text(block.text)
        if not text or is_editorial(text):
            continue
        if paragraph_style(block).startswith("Heading"):
            current_anchor = section["anchors"].get(id(block), "")
        contexts.append((text, current_anchor))
    return contexts


def render_table(
    table: Table,
    prefix: str,
    targets: list[tuple[str, str, re.Pattern]],
    linked: set[str],
) -> str:
    rows_html = []
    for row_index, row in enumerate(table.rows):
        cells_html = []
        seen_cells = set()
        for cell in row.cells:
            identity = id(cell._tc)
            if identity in seen_cells:
                continue
            seen_cells.add(identity)
            tag = "th" if row_index == 0 else "td"
            paragraphs = []
            for paragraph in cell.paragraphs:
                value = clean_text(paragraph.text)
                if value:
                    paragraphs.append(linkify_text(value, prefix, targets, linked))
            content = "<br>".join(paragraphs) if paragraphs else "&nbsp;"
            grid_span = cell._tc.tcPr.gridSpan
            colspan = f' colspan="{grid_span.val}"' if grid_span is not None and int(grid_span.val) > 1 else ""
            cells_html.append(f"<{tag}{colspan}>{content}</{tag}>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return '<div class="table-wrap"><table class="table">' + "".join(rows_html) + "</table></div>"


def render_section_body(
    section: dict,
    doc: Document,
    prefix: str,
    targets: list[tuple[str, str, re.Pattern]],
) -> tuple[str, list[tuple[str, str]]]:
    formats = numbering_formats(doc)
    anchors = section["anchors"]
    linked: set[str] = set()
    output: list[str] = []
    toc: list[tuple[str, str]] = []
    list_kind = None

    def close_list():
        nonlocal list_kind
        if list_kind:
            output.append(f"</{list_kind}>")
            list_kind = None

    for index, block in enumerate(section["blocks"]):
        if isinstance(block, Table):
            close_list()
            output.append(render_table(block, prefix, targets, linked))
            continue
        text = clean_text(block.text)
        style = paragraph_style(block)
        if index == 0 or not text or is_editorial(text):
            continue
        if style.startswith("Heading") or style == "Title":
            close_list()
            level_match = re.search(r"(\d+)$", style)
            level = min(4, max(2, int(level_match.group(1)) if level_match else 2))
            anchor = anchors.get(id(block), slugify(text))
            label = strip_numbering(text)
            if level == 2:
                toc.append((anchor, label))
            css = ' class="section-title"' if level == 2 else ""
            output.append(
                f'<h{level}{css} id="{anchor}">{linkify_text(text, prefix, targets, linked)}</h{level}>'
            )
            continue
        is_list, desired = list_info(block, formats)
        if is_list:
            if list_kind != desired:
                close_list()
                css = "list" if desired == "ul" else "steps"
                output.append(f'<{desired} class="{css}">')
                list_kind = desired
            output.append(f"<li>{linkify_text(text, prefix, targets, linked)}</li>")
            continue
        close_list()
        linked_text = linkify_text(text, prefix, targets, linked)
        if re.search(r"\s=\s", text) and len(text) < 320:
            output.append(f'<div class="formula">{linked_text}</div>')
        else:
            output.append(f"<p>{linked_text}</p>")
    close_list()
    return "\n".join(output), toc


def nav_items(prefix: str, active: str) -> str:
    items = [
        ("inicio", "Início", "index.html"),
        ("modelo", "Modelo de Referência", "modelo.html"),
        ("mapa", "Mapa Geral", "mapa.html"),
        ("modulos", "Módulos", "modulos/index.html"),
        ("dimensoes", "Dimensões", "dimensoes/index.html"),
        ("conceitos", "Conceitos", "conceitos/index.html"),
        ("tecnicas", "Técnicas e Tecnologias", "tecnicas/index.html"),
        ("metodos", "Métodos e Algoritmos", "metodos/index.html"),
        ("artefatos", "Artefatos", "artefatos/index.html"),
        ("metricas", "Métricas", "metricas/index.html"),
        ("problemas", "Problemas Tratados", "problemas/index.html"),
        ("aplicacoes", "Aplicações", "aplicacoes/index.html"),
        ("trilhas", "Trilhas de Estudo", "trilhas/index.html"),
        ("busca", "Busca", "busca.html"),
    ]
    return "".join(
        f'<a{" class=\"active\"" if key == active else ""} href="{prefix}{url}">{label}</a>'
        for key, label, url in items
    )


def shell(
    title: str,
    active: str,
    content: str,
    prefix: str = "",
    breadcrumbs: str = "",
    description: str = "",
) -> str:
    search_url = f"{prefix}busca.html"
    safe_description = html.escape(description or SITE_SUBTITLE, quote=True)
    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <meta name="description" content="{safe_description}">
  <title>{html.escape(title)} · {SITE_TITLE}</title>
  <link rel="stylesheet" href="{prefix}assets/css/main.css">
</head>
<body>
<a class="skip-link" href="#conteudo">Pular para o conteúdo</a>
<div class="app">
  <header class="site-header">
    <div class="header-main">
      <a class="brand" href="{prefix}index.html"><span>UNESP</span>DataLens-RM<small>{SITE_SUBTITLE}</small></a>
      <form class="search" action="{search_url}">
        <label class="sr-only" for="top-search">Pesquisar no conteúdo</label>
        <input id="top-search" name="q" placeholder="Pesquisar conceitos, métodos, métricas..." autocomplete="off">
      </form>
    </div>
    <nav class="top-nav" aria-label="Navegação principal">{nav_items(prefix, active)}</nav>
  </header>
  <main class="main">
    <div class="reading-progress" aria-hidden="true"><span></span></div>
    <article class="content" id="conteudo">{content}</article>
    <footer class="footer"><strong>{SITE_TITLE}</strong> · Engenharia Analítica de Dados com Governança, Qualidade, Proveniência e Reprodutibilidade</footer>
  </main>
</div>
<script src="{prefix}assets/js/main.js"></script>
</body>
</html>
"""


def write_page(relative_path: str, content: str):
    path = ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8", newline="\n")


def hero(eyebrow: str, title: str, description: str, actions: str = "") -> str:
    return f"""<section class="hero">
<span class="badge">{html.escape(eyebrow)}</span>
<h1>{html.escape(title)}</h1>
<p>{html.escape(description)}</p>
{f'<div class="hero-actions">{actions}</div>' if actions else ''}
</section>"""


def card(title: str, description: str, url: str, meta: str = "", number: str = "") -> str:
    marker = f'<span class="module-num">{html.escape(number)}</span>' if number else ""
    return f"""<a class="card card-link" href="{html.escape(url, quote=True)}">
<div class="card-heading">{marker}<div><h3>{html.escape(title)}</h3>{f'<span class="meta">{html.escape(meta)}</span>' if meta else ''}</div></div>
<p>{html.escape(description)}</p><span class="card-action">Explorar <span aria-hidden="true">→</span></span>
</a>"""


def section_kind_and_number(title: str) -> tuple[str, int | None]:
    match = re.search(r"Módulo(?: Transversal)?\s+(\d+)", title, flags=re.IGNORECASE)
    if match:
        return "modulos", int(match.group(1))
    if "Dimensão Complementar" in title:
        return "dimensoes", None
    return "modelo", None


def contexts_for_entry(entry: dict, sections: list[dict], limit: int = 6) -> list[dict]:
    needle = entry["title"].casefold()
    contexts = []
    seen = set()
    for section in sections:
        if not section or section.get("kind") not in {"modulos", "dimensoes", "modelo"}:
            continue
        for text, current_anchor in section["context_paragraphs"]:
            folded = text.casefold()
            if needle not in folded or folded == needle or len(text) < 35:
                continue
            excerpt = text if len(text) <= 300 else text[:297].rstrip() + "…"
            key = excerpt.casefold()
            if key in seen:
                continue
            seen.add(key)
            contexts.append(
                {
                    "excerpt": excerpt,
                    "section": section["display_title"],
                    "url": f"../{section['url']}{'#' + current_anchor if current_anchor else ''}",
                }
            )
            if len(contexts) >= limit:
                return contexts
    return contexts


def build():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    if not SOURCE_PDF.exists():
        raise FileNotFoundError(SOURCE_PDF)
    if not LEXICON_ROOT.exists():
        raise FileNotFoundError(LEXICON_ROOT)

    doc = Document(SOURCE_DOCX)
    entries_by_category = load_entries()
    targets = build_term_targets(entries_by_category)
    sections = split_document(doc)

    module_sections = []
    dimension_sections = []
    overview_section = None
    for section in sections:
        kind, number = section_kind_and_number(section["title"])
        section["kind"] = kind
        section["number"] = number
        section["anchors"] = prepare_heading_anchors(section)
        section["context_paragraphs"] = prepare_context_paragraphs(section)
        if kind == "modulos":
            clean_title = re.sub(r"^Módulo(?: Transversal)?\s+\d+\s*[—-]\s*", "", section["title"])
            section["display_title"] = f"Módulo {number} — {clean_title}"
            section["url"] = f"modulos/m{number:02d}.html"
            module_sections.append(section)
        elif kind == "dimensoes":
            clean_title = re.sub(r"^Dimensão Complementar\s*[—-]\s*", "", section["title"])
            section["display_title"] = clean_title
            section["url"] = f"dimensoes/{slugify(clean_title)}.html"
            dimension_sections.append(section)
        else:
            section["display_title"] = "Modelo de Referência"
            section["url"] = "modelo.html"
            overview_section = section

    module_sections.sort(key=lambda section: section["number"])
    dimension_sections.sort(key=lambda section: section["display_title"].casefold())

    search_records = []

    if overview_section:
        body, toc = render_section_body(overview_section, doc, "", targets)
        toc_html = "".join(f'<a href="#{anchor}">{html.escape(label)}</a>' for anchor, label in toc)
        content = (
            f'<span class="anchor-target" id="{slugify(overview_section["title"])}"></span>'
            + hero(
                "Documento-base",
                f"Modelo de Referência {SITE_TITLE}",
                section_summary(overview_section),
                f'<a class="btn light" href="downloads/{DOWNLOAD_DOCX_NAME}">Baixar DOCX</a>'
                f'<a class="btn ghost" href="downloads/{DOWNLOAD_PDF_NAME}">Baixar PDF</a>',
            )
            + f'<details class="mini-toc" open><summary>Nesta página</summary>{toc_html}</details>'
            + f'<div class="prose module-full">{body}</div>'
        )
        write_page(
            "modelo.html",
            shell(
                "Modelo de Referência",
                "modelo",
                content,
                breadcrumbs='<a href="index.html">Início</a><span>›</span> Modelo de Referência',
                description=section_summary(overview_section),
            ),
        )
        search_records.append(
            {
                "type": "Modelo",
                "title": f"Modelo de Referência {SITE_TITLE}",
                "text": section_summary(overview_section),
                "url": "modelo.html",
            }
        )

    for section in module_sections + dimension_sections:
        prefix = "../"
        body, toc = render_section_body(section, doc, prefix, targets)
        toc_html = "".join(f'<a href="#{anchor}">{html.escape(label)}</a>' for anchor, label in toc)
        if section["kind"] == "modulos":
            eyebrow = f"Módulo {section['number']:02d}"
            active = "modulos"
            parent_label = "Módulos"
            parent_url = "modulos/index.html"
        else:
            eyebrow = "Dimensão complementar"
            active = "dimensoes"
            parent_label = "Dimensões"
            parent_url = "dimensoes/index.html"
        summary = section_summary(section)
        content = (
            f'<span class="anchor-target" id="{slugify(section["title"])}"></span>'
            + hero(eyebrow, section["display_title"], summary)
            + f'<details class="mini-toc"><summary>Sumário desta página · {len(toc)} seções</summary>{toc_html}</details>'
            + f'<div class="prose module-full">{body}</div>'
            + '<div class="page-end"><a href="#conteudo">Voltar ao início da página ↑</a></div>'
        )
        breadcrumbs = (
            f'<a href="../index.html">Início</a><span>›</span>'
            f'<a href="../{parent_url}">{parent_label}</a><span>›</span>{html.escape(section["display_title"])}'
        )
        write_page(
            section["url"],
            shell(
                section["display_title"],
                active,
                content,
                prefix=prefix,
                breadcrumbs=breadcrumbs,
                description=summary,
            ),
        )
        search_records.append(
            {
                "type": "Módulo" if active == "modulos" else "Dimensão",
                "title": section["display_title"],
                "text": summary + " " + " ".join(label for _, label in toc),
                "url": section["url"],
            }
        )

    module_cards = "".join(
        card(
            section["display_title"].split("—", 1)[-1].strip(),
            section_summary(section),
            f"m{section['number']:02d}.html",
            "Pipeline principal" if section["number"] <= 8 else "Capacidade transversal",
            str(section["number"]),
        )
        for section in module_sections
    )
    modules_content = (
        hero(
            "Arquitetura em 16 módulos",
            f"Módulos do {SITE_TITLE}",
            "Os módulos 1 a 8 formam o pipeline técnico-operacional; os módulos 9 a 16 atuam transversalmente sobre todo o ciclo de vida dos dados.",
        )
        + '<div class="legend"><span><i class="dot pipeline"></i> Pipeline principal</span><span><i class="dot transversal"></i> Capacidades transversais</span></div>'
        + f'<section class="grid module-grid">{module_cards}</section>'
    )
    write_page(
        "modulos/index.html",
        shell(
            "Módulos",
            "modulos",
            modules_content,
            prefix="../",
            breadcrumbs='<a href="../index.html">Início</a><span>›</span>Módulos',
        ),
    )

    dimension_cards = "".join(
        card(
            section["display_title"],
            section_summary(section),
            Path(section["url"]).name,
            "Dimensão complementar",
        )
        for section in dimension_sections
    )
    dimensions_content = (
        hero(
            "Perspectivas complementares",
            "Dimensões do modelo",
            "Questões transversais que aprofundam observabilidade, contratos, interoperabilidade, minimização, valor, operação e evolução dos dados.",
        )
        + f'<section class="grid">{dimension_cards}</section>'
    )
    write_page(
        "dimensoes/index.html",
        shell(
            "Dimensões",
            "dimensoes",
            dimensions_content,
            prefix="../",
            breadcrumbs='<a href="../index.html">Início</a><span>›</span>Dimensões',
        ),
    )

    for category, entries in entries_by_category.items():
        label = CATEGORY_LABELS[category]
        description = CATEGORY_DESCRIPTIONS[category]
        cards = []
        for entry in entries:
            excerpt = markdown_excerpt(entry["body"])
            cards.append(card(entry["title"], excerpt, f"{entry['slug']}.html", label.rstrip("s")))
        index_content = (
            hero(f"{len(entries)} verbetes", label, description)
            + '<div class="index-tools"><label for="card-filter">Filtrar nesta categoria</label><input id="card-filter" class="filter-input" type="search" placeholder="Digite um termo..." data-card-filter></div>'
            + f'<section class="grid lexicon-grid">{"".join(cards)}</section>'
        )
        write_page(
            f"{category}/index.html",
            shell(
                label,
                category,
                index_content,
                prefix="../",
                breadcrumbs=f'<a href="../index.html">Início</a><span>›</span>{html.escape(label)}',
                description=description,
            ),
        )
        search_records.append(
            {"type": "Índice", "title": label, "text": description, "url": f"{category}/index.html"}
        )

        for entry in entries:
            article = markdown_to_html(entry["body"])
            contexts = contexts_for_entry(entry, [overview_section] + module_sections + dimension_sections)
            context_html = ""
            if contexts:
                context_cards = "".join(
                    f'<a class="source-context" href="{ctx["url"]}"><span>{html.escape(ctx["section"])}</span><p>{html.escape(ctx["excerpt"])}</p></a>'
                    for ctx in contexts
                )
                context_html = (
                    '<section class="document-context"><h2>No documento de referência</h2>'
                    '<p>Ocorrências selecionadas no conteúdo integral dos módulos e dimensões.</p>'
                    f'<div class="context-list">{context_cards}</div></section>'
                )
            related_method = ""
            if category == "metodos":
                related_method = (
                    f'<div class="box practice"><strong>Classificação editorial.</strong> Este verbete também aparece em '
                    f'<a href="../tecnicas/{entry["slug"]}.html">Técnicas e Tecnologias</a>, pois o documento articula métodos, algoritmos e ferramentas em um mesmo ecossistema.</div>'
                )
            content = (
                hero(label.rstrip("s"), entry["title"], markdown_excerpt(entry["body"]))
                + related_method
                + f'<div class="prose lexicon-entry">{article}</div>'
                + context_html
            )
            breadcrumbs = (
                f'<a href="../index.html">Início</a><span>›</span>'
                f'<a href="index.html">{html.escape(label)}</a><span>›</span>{html.escape(entry["title"])}'
            )
            write_page(
                entry["url"],
                shell(
                    entry["title"],
                    category,
                    content,
                    prefix="../",
                    breadcrumbs=breadcrumbs,
                    description=markdown_excerpt(entry["body"]),
                ),
            )
            search_records.append(
                {
                    "type": label.rstrip("s"),
                    "title": entry["title"],
                    "text": markdown_excerpt(entry["body"]) + " " + " ".join(ctx["excerpt"] for ctx in contexts[:3]),
                    "url": entry["url"],
                }
            )

    first_eight = module_sections[:8]
    transversal = module_sections[8:]
    pipeline_html = "".join(
        f'<a class="flow-node" href="{section["url"]}"><b>{section["number"]}</b><span>{html.escape(section["display_title"].split("—", 1)[-1].strip())}</span></a>'
        for section in first_eight
    )
    transversal_html = "".join(
        f'<a class="layer" href="{section["url"]}"><b>M{section["number"]}</b>{html.escape(section["display_title"].split("—", 1)[-1].strip())}</a>'
        for section in transversal
    )
    dimension_map = "".join(
        f'<a class="dimension-chip" href="{section["url"]}">{html.escape(section["display_title"])}</a>'
        for section in dimension_sections
    )
    map_content = (
        hero(
            "Visão sistêmica",
            "Mapa do conhecimento",
            "Uma leitura compacta das etapas do pipeline, das capacidades transversais e das dimensões complementares.",
        )
        + '<section><h2 class="section-title">Pipeline técnico-operacional</h2>'
        + f'<div class="flow">{pipeline_html}</div></section>'
        + '<section><h2 class="section-title">Capacidades transversais</h2>'
        + f'<div class="layers">{transversal_html}</div></section>'
        + '<section><h2 class="section-title">Dimensões complementares</h2>'
        + f'<div class="dimension-cloud">{dimension_map}</div></section>'
    )
    write_page(
        "mapa.html",
        shell(
            "Mapa Geral",
            "mapa",
            map_content,
            breadcrumbs='<a href="index.html">Início</a><span>›</span>Mapa Geral',
        ),
    )
    search_records.append(
        {
            "type": "Mapa",
            "title": f"Mapa Geral do {SITE_TITLE}",
            "text": "Pipeline principal, capacidades transversais e dimensões complementares.",
            "url": "mapa.html",
        }
    )

    total_entries = sum(len(entries) for entries in entries_by_category.values())
    pipeline_preview = "".join(
        card(
            section["display_title"].split("—", 1)[-1].strip(),
            section_summary(section),
            section["url"],
            "Pipeline" if section["number"] <= 8 else "Transversal",
            str(section["number"]),
        )
        for section in module_sections[:4]
    )
    category_cards = "".join(
        card(
            CATEGORY_LABELS[category],
            CATEGORY_DESCRIPTIONS[category],
            f"{category}/index.html",
            f"{len(entries)} verbetes",
        )
        for category, entries in entries_by_category.items()
    )
    stats = f"""<div class="stats">
<div><strong>{len(module_sections)}</strong><span>módulos</span></div>
<div><strong>{len(dimension_sections)}</strong><span>dimensões</span></div>
<div><strong>{total_entries}</strong><span>verbetes conectados</span></div>
<div><strong>{len(doc.tables)}</strong><span>tabelas do documento</span></div>
</div>"""
    index_content = (
        hero(
            "Modelo de referência",
            "Dados confiáveis, governados e reprodutíveis",
            "Estrutura integrada para projetar, documentar, avaliar e operar pipelines de Engenharia Analítica de Dados — do inventário das fontes aos agentes inteligentes.",
            '<a class="btn light" href="modulos/index.html">Explorar módulos</a><a class="btn ghost" href="mapa.html">Ver mapa geral</a>',
        )
        + stats
        + '<section><div class="section-heading"><div><span class="eyebrow">Comece pelo pipeline</span><h2>Do dado bruto ao ativo analítico</h2></div><a href="modulos/index.html">Ver os 16 módulos →</a></div>'
        + f'<div class="grid module-grid">{pipeline_preview}</div></section>'
        + '<section><div class="section-heading"><div><span class="eyebrow">Enciclopédia conectada</span><h2>Explore por tipo de conhecimento</h2></div></div>'
        + f'<div class="grid">{category_cards}</div></section>'
        + f'<section class="source-banner"><div><span class="eyebrow">Fonte integral</span><h2>Documento de referência completo</h2><p>O conteúdo detalhado dos módulos, dimensões, métricas, artefatos e critérios de validação foi preservado e distribuído em uma arquitetura web interligada.</p></div><div class="source-actions"><a class="btn" href="downloads/{DOWNLOAD_DOCX_NAME}">Baixar DOCX</a><a class="btn secondary" href="downloads/{DOWNLOAD_PDF_NAME}">Baixar PDF</a></div></section>'
    )
    write_page(
        "index.html",
        shell(
            "Início",
            "inicio",
            index_content,
            breadcrumbs="Modelo de referência",
            description=f"{SITE_TITLE}: modelo de referência para Engenharia Analítica de Dados.",
        ),
    )

    search_content = (
        hero(
            "Busca global",
            "Encontre qualquer elemento do modelo",
            "Pesquise módulos, dimensões, conceitos, técnicas, métodos, métricas, artefatos, problemas e aplicações.",
        )
        + '<section class="search-panel"><label for="search-query">O que você procura?</label><div class="search-row"><input id="search-query" type="search" placeholder="Ex.: proveniência, data profiling, completude..." autofocus><button class="btn" id="search-button" type="button">Pesquisar</button></div><div class="search-meta" id="search-meta"></div><div id="search-results"></div></section>'
    )
    search_page = shell(
        "Busca",
        "busca",
        search_content,
        breadcrumbs='<a href="index.html">Início</a><span>›</span>Busca',
    ).replace(
        '<script src="assets/js/main.js"></script>',
        '<script src="assets/js/search-index.js"></script><script src="assets/js/main.js"></script>',
    )
    write_page("busca.html", search_page)

    (ROOT / "assets" / "js").mkdir(parents=True, exist_ok=True)
    search_json = json.dumps(search_records, ensure_ascii=False, separators=(",", ":"))
    (ROOT / "assets" / "js" / "search-index.js").write_text(
        f"window.UNESP_DATALENS_INDEX={search_json};\n", encoding="utf-8", newline="\n"
    )

    downloads = ROOT / "downloads"
    downloads.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_DOCX, downloads / DOWNLOAD_DOCX_NAME)
    shutil.copy2(SOURCE_PDF, downloads / DOWNLOAD_PDF_NAME)
    for legacy_name in (
        "UnespDataLens_Modelo_Referencia_Completo.docx",
        "UnespDataLens-RM_Modelo_Referencia_Completo.docx",
    ):
        legacy_download = downloads / legacy_name
        if legacy_download.exists() and legacy_name != DOWNLOAD_DOCX_NAME:
            legacy_download.unlink()

    readme = f"""# {SITE_TITLE}

Ambiente estático elaborado a partir do documento de referência completo,
com identidade visual e arquitetura de navegação inspiradas no UNESP.IA.

## Abrir

Abra `index.html` em um navegador moderno. O conteúdo não depende de servidor, banco de
dados ou instalação de pacotes.

## Conteúdo

- {len(module_sections)} módulos integrais;
- {len(dimension_sections)} dimensões complementares;
- {total_entries} verbetes conectados;
- busca global local;
- documentos-fonte em DOCX e PDF disponíveis em `downloads/`.

## Reconstruir

O gerador está em `scripts/build_portal.py`. Ele usa o documento-fonte e a taxonomia
editorial local indicada no início do script.
"""
    (ROOT / "README.md").write_text(readme, encoding="utf-8", newline="\n")

    print(
        json.dumps(
            {
                "modules": len(module_sections),
                "dimensions": len(dimension_sections),
                "entries": total_entries,
                "search_records": len(search_records),
                "tables": len(doc.tables),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    build()
